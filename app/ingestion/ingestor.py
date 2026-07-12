import os
import uuid
import logging
from pathlib import Path

import fitz  # PyMuPDF
import pdfplumber
import pytesseract
from PIL import Image
from docx import Document as DocxDocument
import openpyxl

from app.ingestion.chunker import chunk_document
from app.ingestion.embedder import embed_texts
from app.database.connection import SessionLocal
from sqlalchemy import text

logger = logging.getLogger(__name__)

# Tesseract path for Windows
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"


# ── Extractors ────────────────────────────────────────────────────────────────

def extract_pdf(file_path: str) -> str:
    """Extract text from PDF — normal text, tables, and scanned/OCR pages."""
    text_parts = []
    seen = set()

    doc = fitz.open(file_path)

    for page_num, page in enumerate(doc):
        # Step 1 — normal text extraction
        text = page.get_text().strip()

        if text and text not in seen:
            seen.add(text)
            text_parts.append(text)
            logger.info(f"[ingestor] Page {page_num+1} — text extracted ({len(text)} chars)")
        else:
            # Step 2 — no text found, run OCR
            logger.info(f"[ingestor] Page {page_num+1} — no text, running OCR")
            pix = page.get_pixmap(dpi=300)
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
            ocr_text = pytesseract.image_to_string(img).strip()
            if ocr_text and ocr_text not in seen:
                seen.add(ocr_text)
                text_parts.append(ocr_text)
                logger.info(f"[ingestor] Page {page_num+1} — OCR extracted ({len(ocr_text)} chars)")
            else:
                logger.warning(f"[ingestor] Page {page_num+1} — OCR returned nothing")

    # Step 3 — pdfplumber for tables
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    cleaned = [cell.strip() if cell else "" for cell in row]
                    row_text = " | ".join(cleaned)
                    if row_text.strip() and row_text not in seen:
                        seen.add(row_text)
                        text_parts.append(row_text)

    return "\n\n".join(text_parts)


def extract_docx(file_path: str) -> str:
    """Extract text and tables from .docx files."""
    doc = DocxDocument(file_path)
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n\n".join(parts)


def extract_xlsx(file_path: str) -> str:
    """Extract text and numbers from .xlsx files."""
    wb = openpyxl.load_workbook(file_path, data_only=True)
    parts = []

    for sheet in wb.worksheets:
        parts.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cleaned = [str(cell).strip() if cell is not None else "" for cell in row]
            row_text = " | ".join(cleaned)
            if row_text.strip():
                parts.append(row_text)

    return "\n\n".join(parts)


def extract_txt(file_path: str) -> str:
    """Extract text from plain text files."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read()


def extract_content(file_path: str) -> str:
    """Route file to correct extractor based on extension."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return extract_pdf(file_path)
    elif ext == ".docx":
        return extract_docx(file_path)
    elif ext in [".xlsx", ".xls"]:
        return extract_xlsx(file_path)
    elif ext in [".txt", ".md", ".csv"]:
        return extract_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: pdf, docx, xlsx, xls, txt, md, csv")


# ── Main Ingestor ─────────────────────────────────────────────────────────────

def ingest_file(file_path: str, department: str = "general", original_filename: str = None) -> dict:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    filename = original_filename or path.name
    logger.info(f"[ingestor] Starting ingestion: {filename} | department: {department}")

    content = extract_content(file_path)

    if not content.strip():
        raise ValueError(f"No text could be extracted from {filename}")

    logger.info(f"[ingestor] Total characters extracted: {len(content)}")

    chunks = chunk_document(filename, content, department)
    logger.info(f"[ingestor] Chunks created: {len(chunks)}")

    texts = [c["content"] for c in chunks]
    vectors = embed_texts(texts)
    logger.info(f"[ingestor] Vectors created: {len(vectors)}")

    db = SessionLocal()
    doc_id = str(uuid.uuid4())

    try:
        db.execute(text("""
            INSERT INTO documents (id, filename, source, department)
            VALUES (:id, :filename, :source, :department)
        """), {
            "id": doc_id,
            "filename": filename,
            "source": str(file_path),
            "department": department
        })

        for chunk, vector in zip(chunks, vectors):
            chunk_id = str(uuid.uuid4())
            db.execute(text("""
                INSERT INTO chunks
                (id, document_id, content, chunk_index, department, embedding)
                VALUES (:id, :doc_id, :content, :chunk_index, :department, :embedding)
            """), {
                "id": chunk_id,
                "doc_id": doc_id,
                "content": chunk["content"],
                "chunk_index": chunk["chunk_index"],
                "department": department,
                "embedding": str(vector)
            })

        db.commit()
        logger.info(f"[ingestor] SUCCESS — {len(chunks)} chunks saved | Document ID: {doc_id}")

        return {
            "document_id": doc_id,
            "chunks_created": len(chunks)
        }

    except Exception as e:
        db.rollback()
        logger.error(f"[ingestor] ERROR: {e}")
        raise

    finally:
        db.close()


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)

    test_file = "test_document.txt"
    with open(test_file, "w") as f:
        f.write("""Enterprise RAG System Test Document

This document tests the ingestion pipeline.
RAG stands for Retrieval Augmented Generation.
It combines vector search with large language models.
The system retrieves relevant chunks and generates cited answers.
PostgreSQL with pgvector stores 384-dimensional embeddings.
Hybrid retrieval uses semantic search and BM25 with RRF scoring.
ABAC filters ensure department-level access control.
""")

    print("Testing ingestor...")
    result = ingest_file(test_file, department="general")
    print(f"Document ID: {result['document_id']}")
    print(f"Chunks created: {result['chunks_created']}")
    print("Ingestor test PASSED ✅")